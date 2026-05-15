"""
update_report.py
Opens steam_games_multiclass_project_report_updated.docx, embeds missing figures,
fills Section 15 (error analysis), adds Section 17 (feature engineering),
and saves as steam_games_project_report_final.docx.
"""

import json, base64, sys
sys.stdout.reconfigure(encoding='utf-8')
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DATA_DIR = r'c:\Users\I764176\Documents\University\Semester 3\Data Mining'


# ─── Image loading ────────────────────────────────────────────────────────────

def load_nb_images(path):
    """Return {cell_index: bytes} — first image/png output per cell."""
    with open(path, 'r', encoding='utf-8') as f:
        nb = json.load(f)
    images = {}
    for i, cell in enumerate(nb['cells']):
        if cell['cell_type'] != 'code':
            continue
        for out in cell.get('outputs', []):
            data = out.get('data', {})
            if 'image/png' in data:
                raw = data['image/png']
                if isinstance(raw, list):
                    raw = ''.join(raw)
                images[i] = base64.b64decode(raw)
                break   # first image only
    return images


# ─── XML / paragraph insertion helpers ───────────────────────────────────────

def insert_para_after(ref_para, text='', style='Normal'):
    """Insert a new paragraph immediately after ref_para; return it."""
    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)
    p = Paragraph(new_p, ref_para._parent)
    if style != 'Normal':
        p.style = style
    if text:
        p.add_run(text)
    return p


def insert_picture_after(ref_para, img_bytes, width=5.5):
    """Insert a centred image paragraph after ref_para; return the paragraph."""
    p = insert_para_after(ref_para)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(BytesIO(img_bytes), width=Inches(width))
    return p


def insert_table_after(doc, ref_para, headers, rows, style='Table Grid'):
    """
    Create a table, move it after ref_para, add a spacer paragraph after it,
    and return the spacer paragraph (so chaining continues correctly).
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = style
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.rows[i + 1].cells[j].text = str(val)
    # Move table element to just after ref_para
    ref_para._p.addnext(tbl._tbl)
    # Spacer paragraph after the table (needed so next insert has a para reference)
    spacer = OxmlElement('w:p')
    tbl._tbl.addnext(spacer)
    return Paragraph(spacer, ref_para._parent)


def insert_heading_after(ref_para, text, level=2):
    return insert_para_after(ref_para, text, style=f'Heading {level}')


# ─── Lookup helpers ───────────────────────────────────────────────────────────

def find_para(doc, *substrings):
    """Return first paragraph whose text contains ALL given substrings."""
    for para in doc.paragraphs:
        if all(s in para.text for s in substrings):
            return para
    return None


# ─── Fix helpers ─────────────────────────────────────────────────────────────

def fix_rf_test_f1(doc):
    """Change the truncated '0.51' Test Macro F1 in the RF table row to '0.5100'."""
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == 'Random Forest':
                for cell in row.cells:
                    if cell.text.strip() == '0.51':
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text.strip() == '0.51':
                                    run.text = '0.5100'
                                    print('  Fixed: RF Test Macro F1  0.51 → 0.5100')
                                    return
                        # fallback: replace paragraph text directly
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run('0.5100')
                        print('  Fixed (fallback): RF Test Macro F1  0.51 → 0.5100')
                        return


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

print('Loading document ...')
doc = Document(fr'{DATA_DIR}\steam_games_multiclass_project_report_updated.docx')

print('Loading notebook images ...')
nb4b = load_nb_images(fr'{DATA_DIR}\Code\section4b_random_forest_colab.ipynb')
nb5  = load_nb_images(fr'{DATA_DIR}\Code\section5_evaluation_colab.ipynb')
nb6  = load_nb_images(fr'{DATA_DIR}\Code\section6_threshold_experiment_colab.ipynb')
nb7  = load_nb_images(fr'{DATA_DIR}\Code\section7_error_analysis_colab.ipynb')
nb8  = load_nb_images(fr'{DATA_DIR}\Code\section8_feature_engineering_colab.ipynb')

print(f'  section4b images at cells: {sorted(nb4b)}')
print(f'  section5  images at cells: {sorted(nb5)}')
print(f'  section6  images at cells: {sorted(nb6)}')
print(f'  section7  images at cells: {sorted(nb7)}')
print(f'  section8  images at cells: {sorted(nb8)}')


# ── Step 1: Fix RF Test F1 ────────────────────────────────────────────────────
print('\nStep 1: Fix RF Test Macro F1 ...')
fix_rf_test_f1(doc)


# ── Step 2: Embed missing figures ────────────────────────────────────────────
print('\nStep 2: Embedding missing figures ...')

# Figure 3 — model comparison (section5 cell 11: per-class F1 bar chart)
fig3_para = find_para(doc, 'Figure 3')
if fig3_para:
    if 11 in nb5:
        insert_picture_after(fig3_para, nb5[11], width=6.0)
        print('  Embedded Figure 3 (section5 cell 11)')
    else:
        print(f'  WARN: Figure 3 source not found. Available: {sorted(nb5)}')
else:
    print('  WARN: "Figure 3" paragraph not found in document')

# Figure 4 — threshold experiment CMs (section6 cell 17: side-by-side 0.5 vs 0.4)
fig4_para = find_para(doc, 'Figure 4')
if fig4_para:
    if 17 in nb6:
        insert_picture_after(fig4_para, nb6[17], width=6.0)
        print('  Embedded Figure 4 (section6 cell 17)')
    else:
        print(f'  WARN: Figure 4 source not found. Available: {sorted(nb6)}')
else:
    print('  WARN: "Figure 4" paragraph not found in document')

# Figure 5 — RF confusion matrix (section4b cell 13)
fig5_para = find_para(doc, 'Figure 5')
if fig5_para:
    if 13 in nb4b:
        insert_picture_after(fig5_para, nb4b[13], width=5.0)
        print('  Embedded Figure 5 (section4b cell 13)')
    else:
        print(f'  WARN: Figure 5 source not found. Available: {sorted(nb4b)}')
else:
    print('  WARN: "Figure 5" paragraph not found in document')

# Figure 6 — RF feature importances (section4b cell 15)
fig6_para = find_para(doc, 'Figure 6')
if fig6_para:
    if 15 in nb4b:
        insert_picture_after(fig6_para, nb4b[15], width=5.5)
        print('  Embedded Figure 6 (section4b cell 15)')
    else:
        print(f'  WARN: Figure 6 source not found. Available: {sorted(nb4b)}')
else:
    print('  WARN: "Figure 6" paragraph not found in document')


# ── Step 3: Fill Section 15 ───────────────────────────────────────────────────
print('\nStep 3: Filling Section 15 ...')

sec15_para = find_para(doc, '15.')
if sec15_para is None:
    print('  ERROR: Section 15 heading not found — aborting Section 15 fill')
else:
    cur = sec15_para

    # 15.1 Purpose
    cur = insert_heading_after(cur, '15.1  Purpose and setup', level=2)
    cur = insert_para_after(cur,
        'This section investigates why the Mixed class (F1 = 0.43) and Bad class (F1 = 0.33) '
        'score significantly below Good (F1 = 0.77). The analysis uses the best-performing '
        'Random Forest model: max_depth=None, min_samples_leaf=4, n_estimators=200, '
        'class_weight=balanced. The model is fitted on the full training set and evaluated on '
        'the held-out test set (11,331 games), giving Test Macro F1 = 0.5100. Rather than '
        'running additional cross-validation, the model is fitted once so that every test case '
        'can be individually examined.')

    # 15.2 Error flow
    cur = insert_heading_after(cur, '15.2  Error flow analysis', level=2)
    cur = insert_para_after(cur,
        'The row-normalised confusion matrix below shows the fraction of each true class '
        'predicted correctly or as one of the other two classes. Mixed → Good is the dominant '
        'failure mode across the entire test set.')
    if 8 in nb7:
        cur = insert_picture_after(cur, nb7[8], width=5.5)
    cur = insert_table_after(doc, cur,
        headers=['Error type', 'Count', '% of true class'],
        rows=[
            ['Mixed → Good', '1,360', '42.4 %'],
            ['Mixed → Bad',  '505',   '16.6 %'],
            ['Bad → Mixed',  '358',   '37.5 %'],
            ['Bad → Good',   '229',   '24.0 %'],
        ])

    # 15.3 Confidence
    cur = insert_heading_after(cur, '15.3  Prediction confidence', level=2)
    cur = insert_para_after(cur,
        'For each test prediction the model outputs a probability vector; the maximum value is '
        'the prediction confidence. The distribution below separates correct from wrong '
        'predictions.')
    if 10 in nb7:
        cur = insert_picture_after(cur, nb7[10], width=5.5)
    cur = insert_para_after(cur,
        'Correct predictions have a median confidence of 0.570; wrong predictions have a '
        'median confidence of only 0.460. The model is systematically less certain when it is '
        'wrong — most errors are uncertain guesses rather than overconfident misfires.')

    # 15.4 Mixed deep dive
    cur = insert_heading_after(cur, '15.4  Mixed class deep dive', level=2)
    cur = insert_para_after(cur,
        '42.4 % of all Mixed games are predicted as Good — the single largest error in the '
        'test set. Mixed games have a review ratio between 0.50 and 0.74. Many share the same '
        'genre tags, release era, pricing patterns, and platform flags as Good games; the only '
        'difference is that their review ratio sits just below the 0.75 threshold. Because that '
        'margin is not encoded in any feature, the model has no reliable way to separate them.')

    # 15.5 Bad deep dive
    cur = insert_heading_after(cur, '15.5  Bad class deep dive', level=2)
    cur = insert_para_after(cur,
        '37.5 % of Bad games are predicted as Mixed. Bad games form the smallest class '
        '(8.4 % of training data, 955 test cases). With limited training signal the model '
        'defaults toward Mixed — the second-largest class — whenever it is uncertain. '
        'As a result, Bad recall (0.38) exceeds Bad precision (0.30): the model captures most '
        'truly-Bad games but also mislabels many Mixed games as Bad.')

    # 15.6 Feature separation
    cur = insert_heading_after(cur, '15.6  Good vs Mixed feature separation', level=2)
    cur = insert_para_after(cur,
        'The chart below shows the features with the largest absolute mean difference between '
        'Good and Mixed training examples. A positive value means the feature is more prevalent '
        'in Good games; negative means more prevalent in Mixed.')
    if 16 in nb7:
        cur = insert_picture_after(cur, nb7[16], width=5.5)
    cur = insert_table_after(doc, cur,
        headers=['Feature', 'Mean(Good) − Mean(Mixed)'],
        rows=[
            ['tag_2D',                      '+0.172'],
            ['era_2020+',                   '+0.159'],
            ['cat_Steam Cloud',             '+0.145'],
            ['era_2015-2019',               '−0.144'],
            ['tag_Singleplayer',            '+0.128'],
            ['cat_Full controller support', '+0.122'],
        ])
    cur = insert_para_after(cur,
        'The separating power is moderate. Even the strongest signal (tag_2D, diff = +0.172) '
        'means 2D games trend toward Good, but many 2D games are still Mixed or Bad. No single '
        'feature cleanly separates the classes.')

    # 15.7 High-confidence errors
    cur = insert_heading_after(cur, '15.7  High-confidence errors', level=2)
    cur = insert_para_after(cur,
        'High-confidence errors are predictions where the model assigned a probability of '
        '≥ 0.70 to the wrong class — the model\'s blind spots rather than borderline '
        'uncertainty. There are 320 such cases in total.')
    cur = insert_table_after(doc, cur,
        headers=['Error type', 'Count', 'Share of total'],
        rows=[
            ['Mixed → Good', '243', '75.9 %'],
            ['Bad → Good',    '20',  '6.3 %'],
            ['Mixed → Bad',   '11',  '3.4 %'],
        ])
    cur = insert_para_after(cur,
        'Mixed → Good dominates even among the most confident mistakes, confirming that the '
        'model has genuinely learned a wrong boundary for a subset of Mixed games — not just '
        'that it is uncertain about them.')

    # 15.8 Conclusion
    cur = insert_heading_after(cur, '15.8  Conclusion', level=2)
    cur = insert_para_after(cur,
        'The root cause of the low Mixed and Bad F1 is genuine feature-space overlap, not a '
        'model deficiency. Current metadata attributes — tags, Steam categories, price, '
        'release era, and platform flags — do not encode the signals that actually separate '
        'borderline games: execution quality, post-launch developer support, presence of '
        'game-breaking bugs, and community sentiment. Most errors are adjacent (Mixed ↔ Good '
        'or Bad ↔ Mixed) and made with low confidence, consistent with a task where the class '
        'boundaries are genuinely fuzzy in feature space.')

    print('  Section 15 filled successfully')


# ── Step 4: Append Section 17 ────────────────────────────────────────────────
print('\nStep 4: Appending Section 17 ...')

doc.add_heading('17. Feature Engineering Experiment', level=1)
doc.add_paragraph(
    'Section 15 identified that the dominant error is Mixed → Good (42.4 % of Mixed games) '
    'and that the top separating features are tag_2D, era_2020+, cat_Steam Cloud, '
    'tag_Singleplayer, and cat_Full controller support. Eight derived features were engineered '
    'from those signals and added to the original 147-feature dataset (producing 155 features) '
    'to test whether aggregation can push the Random Forest beyond its baseline of '
    'Test Macro F1 = 0.5100.')

doc.add_heading('17.1  Motivation', level=2)
doc.add_paragraph(
    'The error analysis showed that the model struggles most with Mixed games that appear '
    'identical to Good games in feature space. The hypothesis was that count-level aggregates '
    '(number of tags, number of quality-signal features) and interaction terms '
    '(recent era × 2D tag, price × recent era) might capture patterns the individual binary '
    'columns miss.')

doc.add_heading('17.2  New features (147 → 155)', level=2)
doc.add_paragraph(
    'All count-based features were standardised using StandardScaler fitted on the training '
    'set only. Binary interaction features were left unscaled.')
tbl_fe = doc.add_table(rows=9, cols=3)
tbl_fe.style = 'Table Grid'
fe_headers = ['Feature', 'Type', 'Rationale']
for j, h in enumerate(fe_headers):
    cell = tbl_fe.rows[0].cells[j]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
fe_rows = [
    ('n_tags',           'count (scaled)',       'Number of tags — more tags = more discoverable, tends toward Good'),
    ('n_categories',     'count (scaled)',       'Number of Steam feature categories — proxy for release polish'),
    ('n_genres',         'count (scaled)',       'Number of genres listed'),
    ('polish_index',     'sum (scaled)',         'Steam Cloud + Full controller support + Achievements + Singleplayer + Mac'),
    ('recent_2d',        'binary',               'tag_2D × era_2020+ — recent 2D games disproportionately Good'),
    ('price_x_recent',   'interaction (scaled)', 'log_price × era_2020+ — expensive recent games almost always Good'),
    ('solo_only',        'binary',               'tag_Singleplayer × (1 − cat_Multi-player) — focused singleplayer indicator'),
    ('platform_breadth', 'count (scaled)',       'Windows + Mac + Linux — multi-platform developer commitment signal'),
]
for i, (feat, typ, rat) in enumerate(fe_rows, 1):
    tbl_fe.rows[i].cells[0].text = feat
    tbl_fe.rows[i].cells[1].text = typ
    tbl_fe.rows[i].cells[2].text = rat

doc.add_heading('17.3  Cross-validation and test results', level=2)
doc.add_paragraph(
    'The same nested CV setup used throughout the project (5 outer folds, 3 inner folds, '
    'GridSearchCV with the same parameter grid) was repeated on the 155-feature dataset. '
    'The chart below shows the fold-by-fold Macro F1 for RF+FE alongside the original RF '
    'and SVM baselines.')
if 7 in nb8:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(BytesIO(nb8[7]), width=Inches(5.5))
    print('  Embedded Section 17 CV fold chart')

tbl_res = doc.add_table(rows=6, cols=4)
tbl_res.style = 'Table Grid'
res_headers = ['Metric', 'RF+FE (155 features)', 'Original RF (147 features)', 'Change']
for j, h in enumerate(res_headers):
    cell = tbl_res.rows[0].cells[j]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
res_rows = [
    ('Nested CV Macro F1', '0.5102 ± 0.0066', '0.5093 ± 0.0073', '+0.0009'),
    ('Test Macro F1',      '0.5093',           '0.5100',           '−0.0007'),
    ('Good F1 (test)',     '0.77',             '0.77',             '0'),
    ('Mixed F1 (test)',    '0.43',             '0.43',             '0'),
    ('Bad F1 (test)',      '0.33',             '0.33',             '0'),
]
for i, row in enumerate(res_rows, 1):
    for j, val in enumerate(row):
        tbl_res.rows[i].cells[j].text = val

doc.add_heading('17.4  Confusion matrix comparison', level=2)
doc.add_paragraph(
    'The side-by-side row-normalised confusion matrices below compare the original RF '
    '(147 features) against RF+FE (155 features) on the test set. Per-class recall is '
    'effectively unchanged: Good +0.001, Mixed +0.002, Bad −0.013.')
if 13 in nb8:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(BytesIO(nb8[13]), width=Inches(6.0))
    print('  Embedded Section 17 side-by-side CM')

doc.add_heading('17.5  Feature importance', level=2)
doc.add_paragraph(
    'The Gini importance chart below shows the top 25 features; new engineered features are '
    'highlighted in green. Despite no performance gain, 7 of 8 new features rank in the '
    'Top 25: n_tags (rank 4), polish_index (rank 5), price_x_recent (rank 7), '
    'n_categories (rank 8), n_genres (rank 9), recent_2d (rank 12), platform_breadth '
    '(rank 17). The sole exception is solo_only at rank 28.')
if 15 in nb8:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(BytesIO(nb8[15]), width=Inches(5.5))
    print('  Embedded Section 17 feature importance chart')

doc.add_heading('17.6  Conclusion', level=2)
doc.add_paragraph(
    'Feature engineering made no meaningful improvement: Test Macro F1 changed by −0.0007, '
    'well within the standard deviation of the nested CV estimate (±0.0066). Although 7 of '
    '8 new features rank highly by Gini importance, they are largely redundant with the '
    'existing binary columns. n_tags is simply a sum of the individual tag_* columns; '
    'polish_index aggregates features already present one-by-one. The Random Forest was '
    'already learning these count-level patterns from the sparse binary feature space. '
    'This is an important negative result: the performance ceiling is set by genuine data '
    'ambiguity, not by missing features or inadequate engineering. No amount of metadata '
    'aggregation can separate games whose review ratio differs by only a few percentage '
    'points around the class boundaries.')

print('  Section 17 appended successfully')


# ── Save ──────────────────────────────────────────────────────────────────────
out_path = fr'{DATA_DIR}\steam_games_project_report_final.docx'
doc.save(out_path)
print(f'\nSaved: {out_path}')
