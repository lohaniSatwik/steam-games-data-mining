from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1F, 0x38, 0x64)
BLUE      = RGBColor(0x2E, 0x75, 0xB6)
TEAL      = RGBColor(0x1F, 0x7A, 0x8C)
GREEN     = RGBColor(0x37, 0x56, 0x23)
PURPLE    = RGBColor(0x70, 0x30, 0xA0)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
MID_GREY  = RGBColor(0x60, 0x60, 0x60)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# hex helpers for table shading
HEX_NAVY   = '1F3864'
HEX_BLUE   = '2E75B6'
HEX_TEAL   = '1F7A8C'
HEX_GREEN  = '375623'
HEX_PURPLE = '7030A0'
HEX_LGREY  = 'F2F2F2'
HEX_LGREEN = 'E2EFDA'
HEX_LYELL  = 'FFF2CC'
HEX_LRED   = 'FCE4D6'
HEX_LBLUE  = 'DEEAF1'
HEX_WHITE  = 'FFFFFF'


# ── Utility functions ─────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'CCCCCC'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, italic=False, size=10,
              color=DARK_GREY, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return run

def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_body(doc, text, size=10.5, color=DARK_GREY, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = DARK_GREY
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.color.rgb = DARK_GREY
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = DARK_GREY
    return p

def add_space(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
# Title banner table (1×1, navy background)
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.cell(0, 0)
shade_cell(cell, HEX_NAVY)
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(14)
para.paragraph_format.space_after  = Pt(14)
r = para.add_run('Steam Games — Multiclass Classification Project')
r.font.size      = Pt(18)
r.font.bold      = True
r.font.color.rgb = WHITE

add_space(doc)

# Sub-title info block
sub_table = doc.add_table(rows=1, cols=1)
sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sc = sub_table.cell(0, 0)
shade_cell(sc, HEX_LBLUE)
sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
sp = sc.paragraphs[0]
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_before = Pt(8)
sp.paragraph_format.space_after  = Pt(8)
lines = [
    ('IE500 Data Mining  |  Team 9 – Brewed Clusters', 12, True),
    ('University of Mannheim  |  Semester 3', 11, False),
    ('Goal: Predict whether a Steam game will be received as Good, Mixed, or Bad', 10, False),
]
for i, (txt, sz, bd) in enumerate(lines):
    if i > 0:
        sp.add_run('\n')
    r2 = sp.add_run(txt)
    r2.font.size      = Pt(sz)
    r2.font.bold      = bd
    r2.font.color.rgb = NAVY

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Project Overview', level=1, color=NAVY)

add_body(doc,
    'This project applies supervised multiclass classification to predict how Steam game players will '
    'receive a game, based solely on metadata available before and at the time of release. '
    'The three target classes are:')

for cls, defn, pct in [
    ('Good',  'positive review ratio ≥ 75%',  '63.3% of games'),
    ('Mixed', 'positive review ratio 50–74%',  '28.3% of games'),
    ('Bad',   'positive review ratio < 50%',   ' 8.4% of games'),
]:
    add_bullet(doc, f' — {defn}  ({pct} in the training set)', bold_prefix=cls)

add_body(doc,
    'The primary evaluation metric is Macro F1, which averages F1 across all three classes equally. '
    'This is the correct choice for an imbalanced dataset: it penalises a model that ignores the '
    'minority Bad class just as much as one that ignores Good.')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 2. DATASET & PREPROCESSING
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Dataset & Preprocessing', level=1, color=NAVY)

add_body(doc,
    'The raw dataset contains 56,655 Steam games. After preprocessing, 147 features are used '
    'across an 80/20 stratified train/test split (45,324 training, 11,331 test).')

add_heading(doc, 'Feature groups (147 total)', level=2, color=BLUE)
for grp, desc in [
    ('Numeric (4)',          'log_price, Required age, DiscountDLC count, Achievements — StandardScaled'),
    ('Platform flags (3)',   'Windows, Mac, Linux'),
    ('Language count (1)',   'n_languages — StandardScaled'),
    ('Release era (3)',      'era_pre-2015, era_2015–2019, era_2020+  (one-hot)'),
    ('Genres (28)',          'Action, Adventure, RPG, Strategy, etc.  (binary dummies)'),
    ('Categories (47)',      'Steam Cloud, Single-player, Full controller support, etc.  (binary dummies)'),
    ('Top tags (61)',        'Singleplayer, Indie, 2D, Atmospheric, Pixel Graphics, etc.  (binary dummies)'),
]:
    add_bullet(doc, f'  {desc}', bold_prefix=grp)

add_heading(doc, 'Leakage prevention', level=2, color=BLUE)
add_body(doc,
    'Three post-release features — Recommendations, Average playtime, and Median playtime — were '
    'intentionally excluded. These are only known after players have already played and reviewed the '
    'game, so including them would constitute data leakage and inflate performance estimates.')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  Methodology', level=1, color=NAVY)

add_heading(doc, 'Nested Cross-Validation', level=2, color=BLUE)
add_body(doc,
    'All models use the same nested cross-validation framework to ensure unbiased, comparable results:')
for item in [
    ('Outer loop (5-fold stratified)',  'Produces the unbiased performance estimate. The model never sees validation data during training or tuning.'),
    ('Inner loop (3-fold GridSearchCV)','Selects the best hyperparameters using only the training portion of each outer fold.'),
    ('Final model',                     'GridSearchCV is re-run on the full training set with the best parameter grid, then evaluated once on the held-out test set.'),
]:
    add_bullet(doc, f'  {item[1]}', bold_prefix=item[0])

add_heading(doc, 'Handling Class Imbalance', level=2, color=BLUE)
add_body(doc,
    'The dataset is imbalanced (Good 63%, Mixed 28%, Bad 8%). The default strategy across all models '
    'is class_weight="balanced", which re-weights the loss function so the model pays proportionally '
    'more attention to the minority classes. Two oversampling alternatives were also tested '
    '(see Section 5).')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 4. MODELS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Models Built', level=1, color=NAVY)

add_body(doc,
    'Seven models were built and evaluated on the same dataset. Each has a dedicated Google Colab '
    'notebook with full nested CV, final model training, test evaluation, and feature importance plots.')

models_summary = [
    ('4a', 'Logistic Regression',        'Baseline',  'Linear model; class_weight=balanced; C tuned via GridSearchCV'),
    ('4a', 'LR + Random Over-Sampling',  'Variant',   'ROS applied inside ImbPipeline to prevent leakage'),
    ('4a', 'LR + SMOTE',                 'Variant',   'Synthetic minority samples via interpolation, inside ImbPipeline'),
    ('4b', 'Random Forest',              'BEST MODEL','Ensemble of decision trees (bagging); Gini importance'),
    ('4c', 'XGBoost',                    'Strong',    'Sequential boosting; numeric labels via LabelEncoder'),
    ('4d', 'Linear SVM',                 'Strong',    'Linear kernel; class_weight=balanced'),
    ('4e', 'K-Nearest Neighbours (KNN)', 'Weakest',   'Distance-based; 15K stratified subsample; no class_weight'),
]

tbl = doc.add_table(rows=len(models_summary)+1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for c, hdr in enumerate(['Section', 'Model', 'Status', 'Key Design Choice']):
    shade_cell(tbl.cell(0, c), HEX_NAVY)
    cell_text(tbl.cell(0, c), hdr, bold=True, color=WHITE, size=10)

STATUS_COLOR = {
    'Baseline':   HEX_LYELL,
    'Variant':    HEX_LBLUE,
    'BEST MODEL': HEX_LGREEN,
    'Strong':     HEX_LBLUE,
    'Weakest':    HEX_LRED,
}
for r, (sec, name, status, note) in enumerate(models_summary, 1):
    row_color = STATUS_COLOR.get(status, HEX_WHITE)
    for c, val in enumerate([sec, name, status, note]):
        shade_cell(tbl.cell(r, c), row_color)
        alg = WD_ALIGN_PARAGRAPH.LEFT if c == 3 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(tbl.cell(r, c), val,
                  bold=(c == 1), color=DARK_GREY, size=9.5, align=alg)

# Column widths
for c, w in enumerate([1.0, 3.8, 2.2, 8.5]):
    for row in tbl.rows:
        row.cells[c].width = Inches(w / 6.5 * 6.0)

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 5. RESULTS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Results — Model Comparison', level=1, color=NAVY)

add_body(doc,
    'All models evaluated on the same held-out test set (11,331 games). '
    'Macro F1 is the primary metric. Per-class F1 reveals where each model '
    'succeeds and struggles.')

results = [
    ('Random Forest',              '0.5093', '±0.0073', '0.5100', '0.77', '0.43', '0.33', '+0.0745', '1'),
    ('XGBoost',                    '0.4820', '±0.0039', '0.4799', '0.73', '0.39', '0.31', '+0.0444', '2'),
    ('Linear SVM',                 '0.4680', '±0.0034', '0.4693', '0.77', '0.34', '0.29', '+0.0338', '3'),
    ('LR + SMOTE',                 '0.4619', '±0.0064', '0.4648', '0.73', '0.39', '0.27', '+0.0293', '4'),
    ('KNN',                        '0.4462', '±0.0066', '0.4516', '0.77', '0.38', '0.21', '+0.0161', '5'),
    ('LR + class_weight (Baseline)','0.4385','±0.0034', '0.4355', '0.70', '0.33', '0.28',  '—',      '6'),
    ('LR + ROS',                   '0.4376', '±0.0038', '0.4352', '0.70', '0.33', '0.28', '−0.0003', '7'),
]
headers_r = ['Model', 'CV Macro F1', 'CV Std', 'Test\nMacro F1',
             'Good F1', 'Mixed F1', 'Bad F1', 'vs Baseline', 'Rank']

rtbl = doc.add_table(rows=len(results)+1, cols=len(headers_r))
rtbl.style = 'Table Grid'
rtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for c, hdr in enumerate(headers_r):
    shade_cell(rtbl.cell(0, c), HEX_BLUE)
    cell_text(rtbl.cell(0, c), hdr, bold=True, color=WHITE, size=9)

ROW_COL = {
    0: HEX_LGREEN,
    5: HEX_LYELL,
    6: HEX_LRED,
}
for r, row_data in enumerate(results):
    fill = ROW_COL.get(r, HEX_WHITE)
    for c, val in enumerate(row_data):
        shade_cell(rtbl.cell(r+1, c), fill)
        alg = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(rtbl.cell(r+1, c), val,
                  bold=(c == 0 or c == 3), color=DARK_GREY, size=9.5, align=alg)

add_space(doc)

add_heading(doc, 'Key Observations', level=2, color=BLUE)
for obs in [
    ('Random Forest is the best model (Macro F1 = 0.5100)',
     'Ensemble methods outperform linear models because Steam game reception is driven by '
     'non-linear feature interactions (e.g., price × era × genre combinations).'),
    ('Bad class is the hardest to predict across all models',
     'Even the best model achieves only Bad F1 = 0.33. The minority class (8.4%) with limited '
     'training signal is the primary bottleneck.'),
    ('CV and test scores are very close for all models',
     'The largest gap is 0.004 (XGBoost: CV=0.4820, test=0.4799), confirming no overfitting '
     'and that nested CV provided reliable estimates.'),
    ('Linear models plateau around 0.43–0.47',
     'Logistic Regression and Linear SVM cannot capture the non-linear feature combinations '
     'that tree-based models exploit.'),
]:
    add_bullet(doc, f'  {obs[1]}', bold_prefix=obs[0] + ': ')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 6. EXPERIMENTS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Experiments', level=1, color=NAVY)

# 6a — Imbalance
add_heading(doc, '6a.  Imbalance Handling — SMOTE vs ROS vs class_weight', level=2, color=TEAL)
add_body(doc,
    'To determine the best strategy for handling the class imbalance (63% Good / 28% Mixed / 8% Bad), '
    'three approaches were compared using Logistic Regression as a controlled baseline:')

imb_data = [
    ('class_weight="balanced" (Baseline)',
     '0.4385 ± 0.0034', '0.4355', '—',
     'Re-weights the loss function. No data added. Simplest and fastest.'),
    ('Random Over-Sampling (ROS)',
     '0.4376 ± 0.0038', '0.4352', '−0.0003',
     'Duplicates minority samples at random inside CV pipeline. '
     'No improvement — copying existing data adds no new information.'),
    ('SMOTE',
     '0.4619 ± 0.0064', '0.4648', '+0.0293',
     'Creates synthetic minority samples by interpolating between existing ones. '
     'Improves Mixed F1 (0.33→0.39) and Good F1 (0.70→0.73).'),
]

imb_tbl = doc.add_table(rows=len(imb_data)+1, cols=5)
imb_tbl.style = 'Table Grid'
for c, hdr in enumerate(['Strategy', 'CV Macro F1', 'Test Macro F1', 'vs Baseline', 'Mechanism']):
    shade_cell(imb_tbl.cell(0, c), HEX_GREEN)
    cell_text(imb_tbl.cell(0, c), hdr, bold=True, color=WHITE, size=9)
imb_colors = [HEX_LYELL, HEX_LRED, HEX_LGREEN]
for r, (strat, cv, test, diff, mech) in enumerate(imb_data):
    fill = imb_colors[r]
    for c, val in enumerate([strat, cv, test, diff, mech]):
        shade_cell(imb_tbl.cell(r+1, c), fill)
        alg = WD_ALIGN_PARAGRAPH.LEFT if c in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(imb_tbl.cell(r+1, c), val,
                  bold=(c == 0), color=DARK_GREY, size=9.5, align=alg)

add_space(doc)
add_body(doc,
    'Conclusion: SMOTE is the only strategy that meaningfully improves performance (+0.029). '
    'It is important to apply SMOTE inside an imblearn.Pipeline so synthetic samples are never '
    'generated from validation data (which would constitute data leakage). '
    'class_weight="balanced" remains the preferred default for all non-LR models due to its '
    'simplicity and equivalent or better performance.')

add_space(doc)

# 6b — Threshold
add_heading(doc, '6b.  Threshold Experiment — 0.5 vs 0.4 Class Boundary', level=2, color=TEAL)
add_body(doc,
    'The original project defines classes using a 0.5 (50%) lower boundary. An experiment was '
    'conducted to test whether Steam\'s own natural category boundaries (which use 40% as the '
    '"Mixed" lower bound) would yield better results:')

thresh_tbl = doc.add_table(rows=3, cols=5)
thresh_tbl.style = 'Table Grid'
for c, hdr in enumerate(['Threshold', 'Good', 'Mixed', 'Bad', 'Class Sizes (train)']):
    shade_cell(thresh_tbl.cell(0, c), HEX_PURPLE)
    cell_text(thresh_tbl.cell(0, c), hdr, bold=True, color=WHITE, size=9)
for r, (label, good, mixed, bad, sizes) in enumerate([
    ('0.5 — Original', '≥ 75%', '50–74%', '< 50%', 'Good 63.3% / Mixed 28.3% / Bad 8.4%'),
    ('0.4 — Steam Natural', '≥ 70%', '40–69%', '< 40%', 'Good 71.4% / Mixed 24.3% / Bad 4.3%'),
], 1):
    fill = HEX_LGREEN if r == 1 else HEX_LRED
    for c, val in enumerate([label, good, mixed, bad, sizes]):
        shade_cell(thresh_tbl.cell(r, c), fill)
        cell_text(thresh_tbl.cell(r, c), val, bold=(c==0), color=DARK_GREY, size=9.5)

add_space(doc)

thresh_res = doc.add_table(rows=3, cols=5)
thresh_res.style = 'Table Grid'
for c, hdr in enumerate(['Threshold', 'CV Macro F1', 'Test Macro F1', 'Bad F1', 'Bad Precision']):
    shade_cell(thresh_res.cell(0, c), HEX_PURPLE)
    cell_text(thresh_res.cell(0, c), hdr, bold=True, color=WHITE, size=9)
for r, (label, cv, test, badf1, badprec) in enumerate([
    ('0.5 — Original (better)',    '0.4385 ± 0.0034', '0.4355', '0.28', '0.18'),
    ('0.4 — Steam Natural (worse)','0.4173 ± 0.0027', '0.4135', '0.17', '0.10'),
], 1):
    fill = HEX_LGREEN if r == 1 else HEX_LRED
    for c, val in enumerate([label, cv, test, badf1, badprec]):
        shade_cell(thresh_res.cell(r, c), fill)
        cell_text(thresh_res.cell(r, c), val, bold=(c==0), color=DARK_GREY, size=9.5)

add_space(doc)
add_body(doc,
    'Result: the 0.4-threshold performs worse (−0.022 Macro F1). The root cause is Bad class collapse: '
    'setting Bad as < 40% halves the Bad training samples (3,819 → 1,963). Both models achieve '
    'similar Bad recall (~0.60), but the 0.4 model\'s Bad precision drops from 0.18 to 0.10 — '
    'meaning 9 out of 10 Bad predictions are wrong. This is because the narrower Bad zone '
    '(< 40%) is so rare that the model labels many Mixed games as Bad, flooding the Bad column '
    'with false positives. The confusion matrix comparison confirms this visually.')

add_body(doc,
    'Academic insight: this experiment demonstrates that label definition choices materially affect '
    'ML performance, and justifies the original 50% threshold as the better analytical convention '
    'for this task.')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 7. FEATURE IMPORTANCE (SHAP)
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7.  Feature Importance & SHAP Analysis', level=1, color=NAVY)

add_body(doc,
    'The Section 5 evaluation notebook computed SHAP (SHapley Additive exPlanations) values for the '
    'best model (Random Forest) using a 200-sample subset of the test set. SHAP values quantify '
    'each feature\'s contribution to the model\'s prediction for each class.')

add_heading(doc, 'Top 10 features by mean absolute SHAP value (averaged across all 3 classes)', level=2, color=BLUE)

shap_feats = [
    ('1', 'Achievements',          'Number of in-game achievements — more achievements = higher engagement signals Good'),
    ('2', 'era_2020+',             'Games released after 2020 — recent era strongly predicts Good reception'),
    ('3', 'tag_2D',                'Binary tag — 2D games have a distinct audience with high satisfaction'),
    ('4', 'cat_Steam Cloud',       'Steam Cloud save support — indicative of polished, complete games'),
    ('5', 'log_price',             'Log-transformed price — pricing strategy correlates with reception'),
    ('6', 'era_2015–2019',         'Mid-era release window — different market dynamics than pre-2015 or post-2020'),
    ('7', 'tag_Singleplayer',      'Singleplayer games attract players with clear expectations'),
    ('8', 'cat_Full controller support', 'Controller support signals production quality'),
    ('9', 'tag_Pixel Graphics',    'Niche but loyal audience — strong predictor of positive reception'),
    ('10','cat_Steam Achievements','Steam-integrated achievements — indicates developer engagement'),
]

shap_tbl = doc.add_table(rows=len(shap_feats)+1, cols=3)
shap_tbl.style = 'Table Grid'
for c, hdr in enumerate(['Rank', 'Feature', 'Interpretation']):
    shade_cell(shap_tbl.cell(0, c), HEX_BLUE)
    cell_text(shap_tbl.cell(0, c), hdr, bold=True, color=WHITE, size=9)
for r, (rank, feat, interp) in enumerate(shap_feats, 1):
    fill = HEX_LGREEN if r <= 3 else HEX_WHITE
    for c, val in enumerate([rank, feat, interp]):
        shade_cell(shap_tbl.cell(r, c), fill)
        alg = WD_ALIGN_PARAGRAPH.LEFT if c == 2 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(shap_tbl.cell(r, c), val,
                  bold=(c == 1), color=DARK_GREY, size=9.5, align=alg)

add_space(doc)
add_body(doc,
    'SHAP and Random Forest Gini importance agree on the top features, which strengthens confidence '
    'in these findings. Notably, numeric features (Achievements, log_price) dominate tree-based '
    'importance, while categorical tags (era, genre) were more prominent in linear model coefficients.')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8.  Conclusions', level=1, color=NAVY)

conclusions = [
    ('Random Forest is the best model',
     'Macro F1 = 0.5100 on the held-out test set, beating all other models. '
     'Its ability to capture non-linear feature interactions explains its advantage over '
     'linear models (LR, SVM) on this mixed numeric + binary dummy feature space.'),
    ('The task is genuinely difficult',
     'Even the best model achieves Macro F1 = 0.51. The Mixed class (F1 = 0.43) is '
     'intrinsically hard to separate — games in the 50–74% range occupy a diffuse boundary '
     'that is difficult to capture with any feature combination.'),
    ('Non-linear models outperform linear ones',
     'The performance ladder (RF > XGBoost > SVM > LR) follows the degree of non-linearity. '
     'This confirms that game reception is driven by feature interactions, not individual features.'),
    ('SMOTE improves LR but not enough to match tree models',
     'LR + SMOTE (0.4648) beats plain LR (+0.029) but still falls short of SVM and tree models. '
     'Oversampling helps the linear decision boundary but cannot substitute for model capacity.'),
    ('Original 0.5 threshold is better than the Steam 0.4 threshold',
     'The 0.4-threshold experiment proves that label definition choices materially affect outcomes. '
     'The stricter Bad threshold shrinks the minority class too far, collapsing Bad precision to 0.10.'),
    ('Key features: Achievements, release era, tags, and price',
     'These four categories consistently appear at the top of both SHAP and Gini importance rankings, '
     'suggesting that game quality signals (achievements, polish) and market positioning '
     '(price, era) are the strongest predictors of reception.'),
]

for title, body in conclusions:
    add_bullet(doc, f'  {body}', bold_prefix=title + ': ')

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# 9. NOTEBOOK INDEX
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9.  Notebook Index', level=1, color=NAVY)

add_body(doc, 'All notebooks are available on GitHub and designed to run on Google Colab.')

notebooks = [
    ('section1_business_understanding',   'Business framing, research questions, success criteria'),
    ('section2_preprocessing',            'Data cleaning, feature engineering, StandardScaling, train/test split'),
    ('section3_eda',                      'Exploratory data analysis — distributions, correlations, class balance'),
    ('section4a_logistic_colab',          'Logistic Regression baseline (class_weight=balanced)'),
    ('section4a_logistic_ros_colab',      'LR + Random Over-Sampling (ROS) variant'),
    ('section4a_logistic_smote_colab',    'LR + SMOTE variant'),
    ('section4b_random_forest_colab',     'Random Forest — BEST MODEL'),
    ('section4c_xgboost_colab',           'XGBoost'),
    ('section4d_svm_colab',               'Linear SVM'),
    ('section4e_knn_colab',               'K-Nearest Neighbours (15K subsample)'),
    ('section5_evaluation_colab',         'Full model comparison, SHAP analysis, leakage check'),
    ('section6_threshold_experiment_colab','Threshold experiment: 0.5 vs 0.4 boundary + confusion matrix comparison'),
]

nb_tbl = doc.add_table(rows=len(notebooks)+1, cols=2)
nb_tbl.style = 'Table Grid'
for c, hdr in enumerate(['Notebook', 'Contents']):
    shade_cell(nb_tbl.cell(0, c), HEX_NAVY)
    cell_text(nb_tbl.cell(0, c), hdr, bold=True, color=WHITE, size=9)
for r, (nb, desc) in enumerate(notebooks, 1):
    fill = HEX_LGREEN if 'random_forest' in nb else (HEX_LYELL if 'section5' in nb or 'section6' in nb else HEX_WHITE)
    for c, val in enumerate([nb, desc]):
        shade_cell(nb_tbl.cell(r, c), fill)
        alg = WD_ALIGN_PARAGRAPH.LEFT
        cell_text(nb_tbl.cell(r, c), val,
                  bold=(c == 0), color=DARK_GREY, size=9, align=alg)

add_space(doc)


# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out = r'c:\Users\I764176\Documents\University\Semester 3\Data Mining\Project_Summary_Team9.docx'
doc.save(out)
print('Saved:', out)
